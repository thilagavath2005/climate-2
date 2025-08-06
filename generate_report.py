import docx
from docx.shared import Inches

def create_report():
    document = docx.Document()
    document.add_heading('Project Report: Climate Change Impact Prediction using Data Science', 0)

    # 1. Project Title and Abstract
    document.add_heading('1. Project Title and Abstract', level=1)
    p = document.add_paragraph()
    p.add_run('Title: ').bold = True
    p.add_run('Climate Change Impact Prediction using Data Science\n')
    p.add_run('Abstract: ').bold = True
    p.add_run('This project presents a comprehensive data science application designed to predict the impacts of climate change. By leveraging a full-stack web application, the system allows users to upload climate-related data, perform data preprocessing, train multiple machine learning models, and visualize the prediction results. The backend is built with Python (Flask), the frontend with HTML/CSS/JavaScript, and the machine learning pipeline utilizes Scikit-learn and Statsmodels. The primary goal is to provide an accessible and interactive platform for climate data analysis and prediction, enabling stakeholders to better understand and anticipate the effects of climate change.')

    # 2. Introduction
    document.add_heading('2. Introduction', level=1)
    document.add_paragraph().add_run('Problem Statement: ').bold = True
    document.add_paragraph('Climate change is a global challenge with far-reaching consequences. Accurate and timely predictions of its impacts, such as temperature changes and sea-level rise, are crucial for effective planning and mitigation strategies. However, the complexity of climate data and the expertise required for modeling can be significant barriers for many stakeholders.')
    document.add_paragraph().add_run('Objective: ').bold = True
    document.add_paragraph('The main objective of this project is to develop a user-friendly web application that democratizes the process of climate change impact prediction. The system aims to:')
    document.add_paragraph('Provide a seamless workflow for uploading, preprocessing, and analyzing climate data.', style='List Bullet')
    document.add_paragraph('Implement and compare various machine learning models for climate prediction.', style='List Bullet')
    document.add_paragraph('Visualize model performance and prediction results in an intuitive dashboard.', style='List Bullet')
    document.add_paragraph('Offer a flexible and extensible platform for future enhancements.', style='List Bullet')
    document.add_paragraph().add_run('Stakeholders: ').bold = True
    document.add_paragraph('Climate Scientists and Researchers: Can use the platform to quickly test hypotheses and compare model performance on different datasets.', style='List Bullet')
    document.add_paragraph('Policymakers and Government Agencies: Can leverage the predictions to inform environmental policies and disaster management plans.', style='List Bullet')
    document.add_paragraph('Environmental NGOs and Activists: Can use the tool to raise awareness and advocate for climate action.', style='List Bullet')
    document.add_paragraph('Students and Educators: Can use the application as a learning tool to understand the fundamentals of data science and climate modeling.', style='List Bullet')

    # 3. System Overview
    document.add_heading('3. System Overview', level=1)
    document.add_paragraph('The application is designed with a modular architecture, separating the frontend, backend, machine learning pipeline, and data storage into distinct components.')
    document.add_paragraph('Frontend: A web-based user interface built with HTML, CSS, and JavaScript. It allows users to interact with the system, upload data, select options, and view results.', style='List Bullet')
    document.add_paragraph('Backend: A Flask (Python) server that handles user requests, manages the data processing and machine learning workflows, and serves the frontend.', style='List Bullet')
    document.add_paragraph('Machine Learning Pipeline: A set of Python scripts that encapsulate the logic for data preprocessing, model training, evaluation, and prediction.', style='List Bullet')
    document.add_paragraph('Database/File System: The system uses a file-based approach for storing data, models, and results. Raw data, processed data, trained models, and prediction results are stored in a structured directory hierarchy.', style='List Bullet')

    # 4. Technology Stack
    document.add_heading('4. Technology Stack', level=1)
    document.add_paragraph('Frontend:', style='List Bullet')
    document.add_paragraph('HTML5', style='List Bullet 2')
    document.add_paragraph('CSS3', style='List Bullet 2')
    document.add_paragraph('JavaScript (with libraries like Chart.js for visualizations)', style='List Bullet 2')
    document.add_paragraph('Backend:', style='List Bullet')
    document.add_paragraph('Python 3', style='List Bullet 2')
    document.add_paragraph('Flask (for the web framework)', style='List Bullet 2')
    document.add_paragraph('Pandas (for data manipulation)', style='List Bullet 2')
    document.add_paragraph('NumPy (for numerical operations)', style='List Bullet 2')
    document.add_paragraph('Machine Learning:', style='List Bullet')
    document.add_paragraph('Scikit-learn (for Linear Regression, Random Forest, and metrics)', style='List Bullet 2')
    document.add_paragraph('Statsmodels (for the ARIMA model)', style='List Bullet 2')
    document.add_paragraph('Joblib (for saving and loading trained models)', style='List Bullet 2')
    document.add_paragraph('Database:', style='List Bullet')
    document.add_paragraph('File System (for storing CSV datasets, JSON metadata, and serialized models)', style='List Bullet 2')

    # 5. Frontend Features
    document.add_heading('5. Frontend Features', level=1)
    document.add_paragraph('The frontend is designed to guide the user through the entire data science workflow:')
    document.add_paragraph('Dashboard: The main page provides an overview of the system and recent files.', style='List Bullet')
    document.add_paragraph('File Upload: A dedicated page for uploading CSV files containing climate data.', style='List Bullet')
    document.add_paragraph('Data Preprocessing: An interactive interface to select and apply preprocessing steps like handling missing values, scaling, and feature engineering.', style='List Bullet')
    document.add_paragraph('Train/Test Split: A page to split the processed data into training and testing sets, allowing the user to specify the test size and target column.', style='List Bullet')
    document.add_paragraph('Model Training: An interface to select one or more machine learning models to train on the data.', style='List Bullet')
    document.add_paragraph('Predictions and Results: A dashboard to visualize the performance of the trained models (MAE, MSE, RMSE, R-squared) and view the predictions on the test set.', style='List Bullet')
    document.add_paragraph('File Explorer: A utility to browse the project\'s data and results files.', style='List Bullet')

    # 6. Backend Features
    document.add_heading('6. Backend Features', level=1)
    document.add_paragraph('The Flask backend exposes a set of API endpoints and routes to power the frontend:')
    document.add_paragraph('/', style='List Bullet')
    document.add_paragraph('/upload', style='List Bullet')
    document.add_paragraph('/preprocess', style='List Bullet')
    document.add_paragraph('/split-data', style='List Bullet')
    document.add_paragraph('/train-models', style='List Bullet')
    document.add_paragraph('/predictions', style='List Bullet')
    document.add_paragraph('/api/file-content', style='List Bullet')
    document.add_paragraph('/export/<format>', style='List Bullet')

    # 7. ML/AI Models
    document.add_heading('7. ML/AI Models', level=1)
    document.add_paragraph('The project implements three different machine learning models to provide a range of prediction capabilities:')
    document.add_paragraph('Linear Regression: A baseline model that establishes a linear relationship between the input features and the target variable. It\'s fast and easy to interpret.', style='List Bullet')
    document.add_paragraph('Random Forest Regressor: An ensemble model that combines multiple decision trees to improve prediction accuracy and control overfitting. It can capture non-linear relationships and is robust to outliers.', style='List Bullet')
    document.add_paragraph('ARIMA (Autoregressive Integrated Moving Average): A time-series model specifically designed for forecasting. It\'s used when the data has a temporal component, such as predicting future temperature trends.', style='List Bullet')
    document.add_paragraph().add_run('Performance Metrics:').bold = True
    document.add_paragraph('The models are evaluated using the following standard regression metrics:')
    document.add_paragraph('Mean Absolute Error (MAE)', style='List Bullet')
    document.add_paragraph('Mean Squared Error (MSE)', style='List Bullet')
    document.add_paragraph('Root Mean Squared Error (RMSE)', style='List Bullet')
    document.add_paragraph('R-squared (R²)', style='List Bullet')

    # 8. Database Design
    document.add_heading('8. Database Design', level=1)
    document.add_paragraph('The project uses a file-based system for data storage, which is simple and effective for this application\'s scope. The data is organized into the following directories:')
    document.add_paragraph('data/raw/: Stores the original, uploaded CSV files.', style='List Bullet')
    document.add_paragraph('data/processed/: Contains the preprocessed data, ready for model training.', style='List Bullet')
    document.add_paragraph('data/train_test/: Holds the split training and testing datasets.', style='List Bullet')
    document.add_paragraph('data/models/: Stores the serialized, trained machine learning models (using `joblib`).', style='List Bullet')
    document.add_paragraph('data/results/: Contains the prediction results and model performance metrics in JSON format.', style='List Bullet')
    document.add_paragraph('Metadata for each step (upload, preprocessing, splitting, and training) is stored in corresponding JSON files, providing a clear audit trail of the entire process.')

    # 9. Implementation Steps
    document.add_heading('9. Implementation Steps', level=1)
    document.add_paragraph('The project follows a standard data science workflow:')
    document.add_paragraph('Data Collection: The user uploads a CSV file with climate data.', style='List Bullet')
    document.add_paragraph('Data Preprocessing: The DataProcessor class handles:', style='List Bullet')
    document.add_paragraph('Missing Value Imputation: Options to drop rows or fill with mean/median.', style='List Bullet 2')
    document.add_paragraph('Outlier Removal: Using IQR or Z-score methods.', style='List Bullet 2')
    document.add_paragraph('Feature Engineering: Creating interaction terms, polynomial features, and rolling statistics.', style='List Bullet 2')
    document.add_paragraph('Data Scaling: Standardizing or normalizing numeric features.', style='List Bullet 2')
    document.add_paragraph('Train/Test Split: The data is split into training and testing sets to ensure unbiased model evaluation.', style='List Bullet')
    document.add_paragraph('Model Training: The MLPipeline class trains the selected models (Linear Regression, Random Forest, ARIMA) on the training data.', style='List Bullet')
    document.add_paragraph('Model Evaluation: The trained models are evaluated on the test set using the performance metrics.', style='List Bullet')
    document.add_paragraph('Prediction: The best-performing model is used to make predictions on the test data.', style='List Bullet')
    document.add_paragraph('Deployment: The entire application is served via a Flask web server.', style='List Bullet')

    # 10. Challenges & Resolutions
    document.add_heading('10. Challenges & Resolutions', level=1)
    document.add_paragraph('Challenge: Handling diverse and potentially messy climate datasets.', style='List Bullet')
    document.add_paragraph('Resolution: Implemented a robust data preprocessing pipeline with multiple options for handling missing data, outliers, and feature scaling.', style='List Bullet 2')
    document.add_paragraph('Challenge: Integrating different types of machine learning models (regression and time-series).', style='List Bullet')
    document.add_paragraph('Resolution: Designed a flexible MLPipeline class that can accommodate different model types and their specific data requirements.', style='List Bullet 2')
    document.add_paragraph('Challenge: Providing a user-friendly interface for a complex workflow.', style='List Bullet')
    document.add_paragraph('Resolution: Broke down the process into logical steps, with a dedicated page for each stage of the pipeline.', style='List Bullet 2')

    # 11. Screenshots/Results
    document.add_heading('11. Screenshots/Results', level=1)
    document.add_paragraph('(Placeholder for screenshots of the web application\'s UI, such as the dashboard, preprocessing options, and results visualization.)')

    # 12. Conclusion & Future Scope
    document.add_heading('12. Conclusion & Future Scope', level=1)
    document.add_paragraph().add_run('Conclusion: ').bold = True
    document.add_paragraph('This project successfully demonstrates the development of a full-stack data science application for climate change impact prediction. By providing an end-to-end solution, from data upload to model evaluation, it empowers users to perform complex climate data analysis without requiring deep technical expertise.')
    document.add_paragraph().add_run('Future Scope:').bold = True
    document.add_paragraph('Cloud Deployment: Deploy the application to a cloud platform (e.g., AWS, Azure, Google Cloud) for better scalability and accessibility.', style='List Bullet')
    document.add_paragraph('Database Integration: Replace the file-based storage with a robust database system like PostgreSQL or MongoDB to better handle larger datasets and concurrent users.', style='List Bullet')
    document.add_paragraph('Expanded Model Library: Incorporate more advanced machine learning and deep learning models (e.g., Gradient Boosting, LSTMs, Transformers).', style='List Bullet')
    document.add_paragraph('Real-time Data Integration: Connect to real-time climate data APIs to provide up-to-date predictions.', style='List Bullet')
    document.add_paragraph('Enhanced Visualizations: Add more interactive and geographically-based visualizations (e.g., maps).', style='List Bullet')

    document.save('Climate_Forecast_Project_Report.docx')
    print("Report generated successfully.")

if __name__ == '__main__':
    create_report()
